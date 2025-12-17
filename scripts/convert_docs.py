"""Convert Word documents and PDFs to markdown."""
import os
from docx import Document
from PyPDF2 import PdfReader

def docx_to_markdown(input_path, output_path):
    """Convert a Word document to markdown."""
    doc = Document(input_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it looks like a heading (all caps or short with no period)
            if text.isupper() and len(text) < 100:
                lines.append(f"## {text.title()}")
            elif len(text) < 80 and not text.endswith('.') and not text.endswith(':'):
                lines.append(f"### {text}")
            else:
                lines.append(text)
            lines.append("")  # Add blank line after each paragraph

    # Handle tables
    for table in doc.tables:
        lines.append("")
        # Header row
        header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        lines.append(f"| {header} |")
        lines.append("| " + " | ".join("---" for _ in table.rows[0].cells) + " |")
        # Data rows
        for row in table.rows[1:]:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            lines.append(f"| {row_text} |")
        lines.append("")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    print(f"Converted: {input_path} -> {output_path}")

def pdf_to_markdown(input_path, output_path):
    """Convert a PDF to markdown."""
    reader = PdfReader(input_path)
    lines = []

    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text:
            lines.append(f"## Page {page_num}")
            lines.append("")
            lines.append(text)
            lines.append("")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    print(f"Converted: {input_path} -> {output_path}")

if __name__ == "__main__":
    downloads = r"C:\Users\karosebrook\Downloads\Enterpriseprofilebuilder"
    docs_base = r"C:\Users\karosebrook\OneDrive - INT Inc\Desktop\Enterpriseprofilebuilder\public\docs"

    # Convert Word documents
    conversions = [
        (os.path.join(downloads, "INT_Inc_AI_Strategic_Toolkit.docx"),
         os.path.join(docs_base, "reference", "INT_Inc_AI_Strategic_Toolkit.md")),
        (os.path.join(downloads, "INT_Inc_Complete_Personas_Sample.docx"),
         os.path.join(docs_base, "roles", "INT_Inc_Complete_Personas_Sample.md")),
        (os.path.join(downloads, "Next department. at max depth.docx"),
         os.path.join(docs_base, "guides", "Next_Department_Max_Depth.md")),
        (os.path.join(downloads, "Audit-Deep-Dive-Report.md.docx"),
         os.path.join(docs_base, "compliance", "Audit_Deep_Dive_Report.md")),
    ]

    for input_path, output_path in conversions:
        if os.path.exists(input_path):
            try:
                docx_to_markdown(input_path, output_path)
            except Exception as e:
                print(f"Error converting {input_path}: {e}")
        else:
            print(f"File not found: {input_path}")

    # Convert PDF
    pdf_path = os.path.join(downloads, "Expand beyond these tools to others such as automa.pdf")
    pdf_output = os.path.join(docs_base, "guides", "Expand_Beyond_Tools_Automa.md")

    if os.path.exists(pdf_path):
        try:
            pdf_to_markdown(pdf_path, pdf_output)
        except Exception as e:
            print(f"Error converting PDF: {e}")
    else:
        print(f"PDF not found: {pdf_path}")
